from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[1]
OUT_DOCX = ROOT / "docs" / "E-Commerce-Single-Store-Client-Documentation.docx"
ASSETS_DIR = ROOT / "docs" / "assets"


def _set_doc_defaults(doc: Document) -> None:
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")
    style.font.size = Pt(11)


def _title(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(22)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _subtitle(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(12)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def _h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def _p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def _bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def _code_block(doc: Document, lines: list[str]) -> None:
    for line in lines:
        p = doc.add_paragraph()
        run = p.add_run(line)
        run.font.name = "Consolas"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Consolas")
        run.font.size = Pt(10)


def _img(doc: Document, path: Path, width_in: float, caption: str | None = None) -> None:
    if not path.exists():
        doc.add_paragraph(f"[Missing image: {path.name}]")
        return
    doc.add_picture(str(path), width=Inches(width_in))
    if caption:
        p = doc.add_paragraph(caption)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def build_doc() -> Document:
    doc = Document()
    _set_doc_defaults(doc)

    today = date.today().isoformat()

    _title(doc, "E‑Commerce Single Store")
    _subtitle(doc, "Client Documentation (Features, Admin Guide, Setup & Operations)")
    _subtitle(doc, f"Generated on: {today}")
    doc.add_page_break()

    _h1(doc, "1. Overview")
    _p(
        doc,
        "This project is a single‑store e‑commerce system consisting of a customer website, an admin dashboard, and a backend API.",
    )
    _bullets(
        doc,
        [
            "Storefront: product browsing, cart, checkout, order tracking, user account features.",
            "Admin panel: manage products, categories, banners, orders, shipping zones, payment methods, coupons, reports, and website settings.",
            "Backend: Node.js/Express API with MongoDB persistence and optional integrations (email, courier, analytics).",
        ],
    )

    _img(
        doc,
        ASSETS_DIR / "architecture.png",
        width_in=6.8,
        caption="Figure 1 — System Architecture (Single Store).",
    )

    _h1(doc, "2. Customer‑Facing Features (Storefront)")
    _bullets(
        doc,
        [
            "Product catalog: browse products, categories, and brands.",
            "Product details: images, descriptions, and reviews (where enabled).",
            "Cart & checkout: add to cart, update quantities, and place orders.",
            "Shipping calculation: shipping costs are controlled by admin‑configured shipping zones.",
            "Coupons/discounts: customers can apply coupon codes at checkout (if configured).",
            "Order tracking: track an order from the public tracking page.",
            "Customer account: registration/login, saved addresses, order history, wishlist.",
            "Pages: about, contact, FAQs, policy pages, and marketing landing pages by slug.",
        ],
    )

    _h1(doc, "3. Admin Panel (Dashboard) — What You Can Manage")
    _p(doc, "Admins access the dashboard at `/dashboard` after logging in.")
    _img(
        doc,
        ASSETS_DIR / "admin-modules.png",
        width_in=6.8,
        caption="Figure 2 — Admin Panel Modules (Illustration).",
    )
    _h2(doc, "3.1 Main Admin Modules")
    _bullets(
        doc,
        [
            "Website Setup: store branding, theme settings, pages, landing pages, and policies.",
            "Catalog: products, categories, brands, and homepage banners.",
            "Orders: view/update orders, add manual orders, and (optionally) courier/consignment tools.",
            "Commerce: shipping zones, payment methods, coupons.",
            "Operations: inventory center, abandoned orders, customer risk controls, support tickets.",
            "Insights: business reports, SEO & analytics dashboards.",
            "Security: admin users, roles/permissions (super admin can control restricted modules).",
        ],
    )

    _h1(doc, "4. Order Lifecycle (How Orders Flow)")
    _img(
        doc,
        ASSETS_DIR / "order-lifecycle.png",
        width_in=6.8,
        caption="Figure 3 — Typical Order Lifecycle.",
    )

    _h1(doc, "5. Environment & Prerequisites")
    _bullets(
        doc,
        [
            "Node.js (backend + frontend).",
            "MongoDB (database).",
            "Cloudinary account (recommended for image hosting).",
            "SMTP email account (recommended for password reset and store emails).",
        ],
    )

    _h1(doc, "6. Quick Start (Local Development)")
    _h2(doc, "6.1 Backend setup")
    _p(doc, "From the repository root:")
    _code_block(
        doc,
        [
            "cd backend",
            "copy .env.example .env   (Windows PowerShell: Copy-Item .env.example .env)",
            "npm install",
            "npm run dev",
        ],
    )
    _p(
        doc,
        "Backend default port: `5000` (configurable via `backend/.env`).",
    )

    _h2(doc, "6.2 Frontend setup")
    _code_block(
        doc,
        [
            "cd frontend",
            "copy .env.example .env   (Windows PowerShell: Copy-Item .env.example .env)",
            "npm install",
            "npm run dev",
        ],
    )
    _p(doc, "Frontend default dev URL: `http://localhost:5173`.")

    _h2(doc, "6.3 MongoDB setup")
    _bullets(
        doc,
        [
            "Install MongoDB locally or use a managed MongoDB cluster.",
            "Set `MONGODB_URI` in `backend/.env` (example: `mongodb://localhost:27017/ecommerce_office`).",
        ],
    )

    _h1(doc, "7. Production Notes (Deployment Checklist)")
    _bullets(
        doc,
        [
            "Set strong `JWT_SECRET` in production.",
            "Set `NODE_ENV=production` on the backend.",
            "Update `FRONTEND_URL` in `backend/.env` to your deployed frontend URL.",
            "Update `VITE_API_URL` in `frontend/.env` to your deployed backend API URL (ending with `/api`).",
            "Ensure MongoDB backups are configured.",
            "Ensure uploads/media are persisted (Cloudinary recommended).",
        ],
    )

    _h1(doc, "8. First‑Time Admin Setup (Required)")
    _p(
        doc,
        "Important: the first user registered in a fresh database becomes the primary Admin (and Super Admin).",
    )
    _bullets(
        doc,
        [
            "Start backend + frontend with an empty MongoDB database.",
            "Open the website and register your first account (name/email/phone/password).",
            "Log in and open `/dashboard` to access Admin modules.",
            "Immediately update environment secrets and keep admin credentials private.",
        ],
    )

    _h1(doc, "9. Operating the Store (Admin How‑To)")
    _p(
        doc,
        "This section explains the recommended day‑to‑day workflows inside the Admin Panel. "
        "Exact labels can differ slightly depending on your theme and enabled modules.",
    )
    _h2(doc, "9.1 Website setup")
    _bullets(
        doc,
        [
            "Set store name, contact details, and theme/branding settings.",
            "Configure homepage banners and highlight content.",
            "Create landing pages for marketing campaigns (URL: `/lp/<slug>`).",
            "Maintain policy pages (URL: `/policy/<policyType>`).",
        ],
    )
    _p(
        doc,
        "Recommended first steps (in order): set branding → add banners → add policies → create landing pages.",
    )

    _h2(doc, "9.2 Catalog management")
    _bullets(
        doc,
        [
            "Create categories and brands first, then add products.",
            "For each product: title, price/sale price, stock, images, category, and description.",
            "Optional: add SKU, low-stock threshold, and product variations (if used by your store).",
            "Use the modify product module for editing, disabling, or removing items from sale.",
        ],
    )
    _bullets(
        doc,
        [
            "Suggested workflow: Categories → Brands → Products → Review the storefront product page.",
            "Tip: keep product titles consistent and use clear photos for better conversion.",
        ],
    )

    _h2(doc, "9.3 Shipping zones")
    _bullets(
        doc,
        [
            "Create shipping zones to control delivery costs by region.",
            "Review shipping prices periodically and update based on courier rates.",
        ],
    )
    _p(
        doc,
        "If shipping looks incorrect during checkout, re-check shipping zones first.",
    )

    _h2(doc, "9.4 Payment methods")
    _bullets(
        doc,
        [
            "Enable and configure available payment methods (e.g., Cash on Delivery, manual bank/bKash instructions).",
            "Confirm checkout and order confirmation behavior using a test order.",
        ],
    )
    _p(
        doc,
        "Payments are customer-facing: always place at least one test order after changes.",
    )

    _h2(doc, "9.5 Orders (daily workflow)")
    _bullets(
        doc,
        [
            "Open Orders and review new orders.",
            "Confirm customer information, address, and items.",
            "Update order status as you process (processing → shipped → delivered, depending on policy).",
            "Use the Add Order module to place an order manually when needed (phone/in‑store orders).",
        ],
    )
    _bullets(
        doc,
        [
            "Recommended daily routine: check new orders → confirm stock → pack → dispatch → sync tracking → close tickets.",
            "If a customer asks for changes, update the order before shipping and keep an internal note (if supported).",
        ],
    )

    _h2(doc, "9.6 Courier integration (optional)")
    _p(
        doc,
        "If courier integration is enabled and configured, the admin can generate consignments, print labels, and sync tracking updates from the order details view.",
    )
    _bullets(
        doc,
        [
            "Open Admin → Courier Settings and enter courier credentials.",
            "Open an order, generate a courier consignment, then print/download label (if supported).",
            "Use “Sync Status” to pull the latest tracking state into the order record.",
        ],
    )

    _h2(doc, "9.7 Reports, SEO & analytics")
    _bullets(
        doc,
        [
            "Use Business Reports for sales trends, order summaries, and performance snapshots.",
            "Use SEO & Analytics to monitor page performance and marketing impact.",
            "Follow the Google Analytics setup guide if you want GA tracking enabled.",
        ],
    )

    _h2(doc, "9.8 Support & customer risk")
    _bullets(
        doc,
        [
            "Support Tickets: track customer issues and close them after resolution.",
            "Customer Risk: review suspicious customers, blacklist abusive buyers if needed.",
            "Abandoned Orders: review incomplete checkouts and follow up (process-based).",
        ],
    )

    _h2(doc, "9.9 Admin users & permissions")
    _p(
        doc,
        "The first registered user in a new database becomes Admin and also Super Admin. "
        "Super Admin can control restricted modules and permissions.",
    )
    _bullets(
        doc,
        [
            "Create separate admin accounts for staff (do not share one password).",
            "Grant only the modules each staff member needs (least privilege).",
            "Remove access promptly when staff roles change.",
        ],
    )

    _h1(doc, "10. Security, Backup & Maintenance")
    _h2(doc, "10.1 Security essentials")
    _bullets(
        doc,
        [
            "Use a long, unique `JWT_SECRET` in production.",
            "Use strong admin passwords (8+ characters minimum; longer recommended).",
            "Keep server access (SSH / hosting panel) limited to trusted operators.",
            "Prefer Cloudinary (or durable object storage) for images to avoid losing uploads on redeploy.",
        ],
    )
    _h2(doc, "10.2 Backups")
    _bullets(
        doc,
        [
            "MongoDB: schedule automated backups (daily recommended).",
            "Media: ensure Cloudinary account access is secured; export critical assets if required by policy.",
            "Config: keep `.env` values in a secure password manager (never commit secrets).",
        ],
    )
    _h2(doc, "10.3 Routine maintenance")
    _bullets(
        doc,
        [
            "Review order/report dashboards weekly to spot issues early.",
            "Check email deliverability monthly (password resets, notifications).",
            "Keep Node.js and dependencies updated during scheduled maintenance windows.",
        ],
    )

    _h1(doc, "11. Troubleshooting (Common Issues)")
    _h2(doc, "11.1 API not reachable from the website")
    _bullets(
        doc,
        [
            "Check `frontend/.env` → `VITE_API_URL` points to the backend `/api` base URL.",
            "Check backend CORS: `backend/.env` → `FRONTEND_URL` matches the deployed frontend URL.",
            "Confirm backend is running and reachable on its port (default 5000).",
        ],
    )
    _h2(doc, "11.2 Emails not sending")
    _bullets(
        doc,
        [
            "Verify SMTP credentials in `backend/.env` (`EMAIL_HOST`, `EMAIL_PORT`, `EMAIL_USER`, `EMAIL_PASS`).",
            "If using Gmail, use an App Password (not your normal Gmail password).",
            "Check server logs for SMTP authentication errors.",
        ],
    )
    _h2(doc, "11.3 Image upload issues")
    _bullets(
        doc,
        [
            "Confirm Cloudinary keys are set in `backend/.env`.",
            "If using local uploads, ensure the backend `uploads/` directory is writable and persisted.",
        ],
    )
    _h2(doc, "11.4 Courier integration issues (optional)")
    _bullets(
        doc,
        [
            "Re-check courier credentials and provider settings in Admin → Courier Settings.",
            "Use the courier troubleshooting FAQ documents for provider-specific error codes.",
        ],
    )

    _h1(doc, "12. Appendix: Key Configuration Variables")
    _p(doc, "Backend (`backend/.env`) highlights:")
    _code_block(
        doc,
        [
            "PORT=5000",
            "NODE_ENV=development",
            "MONGODB_URI=mongodb://localhost:27017/ecommerce_office",
            "JWT_SECRET=replace-with-secure-secret",
            "FRONTEND_URL=http://localhost:5173",
            "CLOUDINARY_CLOUD_NAME=...",
            "CLOUDINARY_API_KEY=...",
            "CLOUDINARY_API_SECRET=...",
            "EMAIL_HOST=smtp.gmail.com",
            "EMAIL_PORT=587",
            "EMAIL_USER=...",
            "EMAIL_PASS=...  (Gmail App Password recommended)",
        ],
    )
    _p(doc, "Frontend (`frontend/.env`) highlights:")
    _code_block(doc, ["VITE_API_URL=http://localhost:5000/api"])

    _h1(doc, "13. Additional Project Guides")
    _bullets(
        doc,
        [
            "Courier integration: `COURIER_QUICK_SETUP.md`, `COURIER_INTEGRATION_GUIDE.md`, `COURIER_TROUBLESHOOTING_FAQ.md`",
            "Google Analytics: `GOOGLE_ANALYTICS_SETUP.md`",
        ],
    )

    return doc


def main() -> None:
    doc = build_doc()
    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print(f"Wrote: {OUT_DOCX}")


if __name__ == "__main__":
    main()
